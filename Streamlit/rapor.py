from docx import Document
import io
import tempfile
import siniflandirma

def create_rapor(model,model_name, X_train, X_test, y_train, y_test, y_pred):
    # Yeni bir Word belgesi oluştur
    doc = Document()
    doc.add_heading(f'{model_name} Raporu', 0)

    # Model sonuçlarını ekle
    doc.add_heading('Model Sonuçları:', level=1)
    accuracy = accuracy_score(y_test, y_pred)
    precision = precision_score(y_test, y_pred, average='weighted')
    recall = recall_score(y_test, y_pred, average='weighted')
    f1 = f1_score(y_test, y_pred, average='weighted')
    
    doc.add_paragraph(f"Doğruluk (Accuracy): {accuracy:.2f}")
    doc.add_paragraph(f"Kesinlik (Precision): {precision:.2f}")
    doc.add_paragraph(f"Duyarlılık (Recall): {recall:.2f}")
    doc.add_paragraph(f"F1 Skoru: {f1:.2f}")

    # Confusion Matrix görseli ekle
    cm = confusion_matrix(y_test, y_pred)
    fig, ax = plt.subplots(figsize=(8, 6))
    sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", cbar=False, ax=ax)
    
    # Geçici dosya oluştur ve görseli kaydet
    with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
        cm_image_path = tmpfile.name
        fig.savefig(cm_image_path)
    
    doc.add_paragraph('Confusion Matrix:')
    doc.add_picture(cm_image_path)

    # ROC-AUC görseli ekle
    if hasattr(model, "predict_proba"):
        y_prob = model.predict_proba(X_test)[:, 1]
        fpr, tpr, _ = roc_curve(y_test, y_prob)
        roc_auc = auc(fpr, tpr)
        fig, ax = plt.subplots(figsize=(8, 6))
        ax.plot(fpr, tpr, label=f'ROC Curve (AUC = {roc_auc:.2f})')
        ax.plot([0, 1], [0, 1], 'k--')
        ax.set_xlabel("False Positive Rate")
        ax.set_ylabel("True Positive Rate")
        ax.legend()
        
        # Geçici dosya oluştur ve ROC-AUC görselini kaydet
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmpfile:
            roc_image_path = tmpfile.name
            fig.savefig(roc_image_path)
        
        doc.add_paragraph('ROC-AUC Eğrisi:')
        doc.add_picture(roc_image_path)

    # Classification Report ekle
    doc.add_heading('Sınıflandırma Raporu:', level=1)
    report = classification_report(y_test, y_pred)
    doc.add_paragraph(report)

    # Dosyayı kaydet
    rapor_path = "C:/Users/kardelen.erdem/Desktop/PROJELER/AutoML/siniflandirma_raporu.docx"
    doc.save(rapor_path)

    return rapor_path
