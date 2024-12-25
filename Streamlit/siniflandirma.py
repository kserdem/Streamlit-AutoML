import streamlit as st
import pandas as pd
from datetime import datetime
from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.svm import SVC
from sklearn.neighbors import KNeighborsClassifier
from sklearn.metrics import accuracy_score, precision_score, recall_score, f1_score
from sklearn.preprocessing import StandardScaler, LabelEncoder
from sklearn.metrics import confusion_matrix
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.metrics import roc_curve, auc
from sklearn.metrics import classification_report
from docx import Document
import io
import tempfile
import rapor

def siniflandirma_sayfasi():

    ####
    # Rapor oluşturma fonksiyonu
    def create_rapor(model_name, X_train, X_test, y_train, y_test, y_pred,
                     selected_features=None,  feature_importances=None, params=None, feature_names=None):
        # Word belgesi oluşturuluyor
        doc = Document()
        doc.add_heading(f'{model_name} Modeli Raporu', 0)

        # Eğitim ve test verisi hakkında bilgi ekleniyor
        doc.add_heading('Veri Bilgisi', level=1)
        doc.add_paragraph(f'Eğitim Verisi Boyutu: {X_train.shape}')
        doc.add_paragraph(f'Test Verisi Boyutu: {X_test.shape}')

        # Modelin performansı hakkında bilgi ekleniyor
        doc.add_heading('Model Performansı', level=1)
        doc.add_paragraph(f"Doğruluk (Accuracy): {accuracy_score(y_test, y_pred)}")
        doc.add_paragraph(f"Kesinlik (Precision): {precision_score(y_test, y_pred, average='weighted')}")
        doc.add_paragraph(f"Duyarlılık (Recall): {recall_score(y_test, y_pred, average='weighted')}")
        doc.add_paragraph(f"F1 Skoru: {f1_score(y_test, y_pred, average='weighted')}")


        # Özellik seçimi bilgisi
        if selected_features:
            doc.add_heading('Seçilen Özellikler', level=1)
            doc.add_paragraph(f'Seçilen Özellikler: {", ".join(selected_features)}')
        else:
            doc.add_paragraph('Model eğitiminde veri setindeki tüm değişkenler kullanıldı.')

        #feature importance
        if hasattr(model_name, 'feature_importances_') and feature_importances is not None:
            doc.add_heading('Özellik Önemleri', level=1)
            sorted_importances = sorted(zip(feature_names, feature_importances), key=lambda x: x[1], reverse=True)
            for feature, importance in sorted_importances:
                doc.add_paragraph(f'{feature}: {importance:.4f}')

        # Model hiperparametreleri
        if params:
            doc.add_heading('Model Hiperparametreleri', level=1)
            for param, value in params.items():
                doc.add_paragraph(f'{param}: {value}')

        # Performans metrikleri
        doc.add_heading('Model Performansı', level=1)
        doc.add_paragraph(f"Doğruluk (Accuracy): {accuracy_score(y_test, y_pred):.4f}")
        doc.add_paragraph(f"Kesinlik (Precision): {precision_score(y_test, y_pred, average='weighted'):.4f}")
        doc.add_paragraph(f"Duyarlılık (Recall): {recall_score(y_test, y_pred, average='weighted'):.4f}")
        doc.add_paragraph(f"F1 Skoru: {f1_score(y_test, y_pred, average='weighted'):.4f}")

        # Sınıflandırma matrisi
        doc.add_heading('Sınıflandırma Matrisi', level=1)
        report = classification_report(y_test, y_pred, output_dict=False)
        doc.add_paragraph(report)

        # Confusion Matrix
        doc.add_heading('Karışıklık Matrisi', level=1)
        cm = confusion_matrix(y_test, y_pred)
        cm_str = "\n".join(["\t".join(map(str, row)) for row in cm])
        doc.add_paragraph(f"Karışıklık Matrisi:\n{cm_str}")

        # Raporu kaydediyoruz
        rapor_path = f'{model_name}_raporu_{datetime.now().strftime("%Y%m%d%H%M%S")}.docx'
        doc.save(rapor_path)
        return rapor_path


    # Başlık
    st.title("CATML Sınıflandırma Arayüzü")

    # Veri kümesi yükleme

    datasets = {
    "Heart Dataset": "cleaned_merged_heart_dataset.csv",
    "Drug 200 Dataset": "drug200.csv",
    "Heart Dataset Metadata": "heart_metadata.csv",
    "Drug 200 Dataset Metadata": "drug200_metadata.csv"
    }

    # Kullanıcıdan veri kümesi ve metadata seçmesini isteme
    dataset_choice = st.selectbox("Veri Kümesini Seçin", ["Heart Dataset", "Drug 200 Dataset"])
    #metadata_choice = st.selectbox("Metadata Seçin", ["Heart Dataset Metadata", "Drug 200 Dataset Metadata"])

    # Seçime göre veri yükleme
    data = pd.read_csv(datasets[dataset_choice])
    metadata = pd.read_csv(datasets[dataset_choice+ " Metadata"])



    #uploaded_file = st.file_uploader("Veri Kümesini Yükleyin (CSV formatında)", type="csv")
    #uploaded_metadata = st.file_uploader("Metadata Yükleyin (... formatında)", type="csv")
    if dataset_choice:
        data = pd.read_csv(datasets[dataset_choice])
        #metadata = pd.read_csv(datasets[metadata_choice])
        
        tab1, tab2, tab3, tab4 = st.tabs(["Veri Kümesi", "Tanımlayıcı İstatistikler", "Veri Seti Özeti", "Değişken Tanımlamaları"])
        with tab1:
            with st.expander("Veri Kümesinin İlk 5 Satırı"):
                st.write(data.head())

        with tab2:
            st.write(data.describe().T)

        with tab3:
            info = pd.DataFrame({
                'Veri Tipi': data.dtypes,
                'Eksik Değerler': data.isnull().sum(),
                'Toplam Değer': data.count(),
                'Benzersiz Değerler': data.nunique()})
            st.table(info)

        with tab4:
            st.write(metadata.T)

        # Ön işleme seçenekleri
        st.sidebar.header("Veri Ön İşleme Adımları")
        missing_option = st.sidebar.selectbox("Eksik Değer İşleme", ("Kaldır", "Ortalama ile Doldur", "Medyan ile Doldur"))
        scaling_option = st.sidebar.checkbox("Özellik Ölçekleme (StandardScaler)")
        encoding_option = st.sidebar.checkbox("Kategorik Değişkenleri Kodla (Label Encoding)")

        # Eksik değer işleme
        if missing_option == "Kaldır":
            data = data.dropna()
        elif missing_option == "Ortalama ile Doldur":
            data = data.fillna(data.mean())
        elif missing_option == "Medyan ile Doldur":
            data = data.fillna(data.median())

        # Kategorik değişkenleri kodlama
        if encoding_option:
            label_encoders = {}
            for column in data.select_dtypes(include=['object']).columns:
                le = LabelEncoder()
                data[column] = le.fit_transform(data[column])
                label_encoders[column] = le

        # Özellik ölçekleme
        if scaling_option:
            scaler = StandardScaler()
            numeric_columns = data.select_dtypes(include=['float64', 'int64']).columns
            data[numeric_columns] = scaler.fit_transform(data[numeric_columns])

        # Özellik ve hedef değişkenleri seçme
        target = st.selectbox("Hedef Değişkeni Seçin", data.columns)
        X = data.drop(columns=[target])
        y = data[target]

        # Veri kümesini ayırma
        def test_train(X, y, params=None):
            if params is None:
                params = {}

            params["test_size"] = st.sidebar.slider("Test Oranını Belirleyin", 0.2, 0.5, 0.3)  # Default değer 0.3
            X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=params["test_size"], random_state=42)
            return X_train, X_test, y_train, y_test
        X_train, X_test, y_train, y_test = test_train(X, y)

        # Model seçimi
        model_name = st.sidebar.selectbox("Model", ("Lojistik Regresyon", "Rasgele Orman", "SVM", "K-En Yakın Komşu"))

        # Hiperparametre ayarları
        def model_hiperparametreleri(model_name):
            params = {}
            if model_name == "Lojistik Regresyon":
                
                params["max_iter"] = st.sidebar.slider("Maksimum İterasyon", 100, 500, 200)
                params["C"] = st.sidebar.slider("C (Regularization Strength)", 0.01, 10.0, 1.0)
                params["solver"] = st.sidebar.selectbox("Solver (Çözümleyici)", ('lbfgs', 'liblinear', 'newton-cg','newton-cholesky','sag','saga'))
                if params["solver"] =='lbfgs':
                    penalty_options = ['l2', None]
                elif params["solver"] =='liblinear':
                    penalty_options = ['l1', 'l2']
                elif params["solver"] =='newton-cg':
                    penalty_options = ['l2', None]
                elif params["solver"] =='newton-cholesky':
                    penalty_options = ['l2', None]
                elif params["solver"] =='sag':
                    penalty_options = ['l2', None]
                elif params["solver"] =='saga':
                    penalty_options = ['l1', 'l2', None]
                params["penalty"] = st.sidebar.selectbox("Penalty (Ceza Parametresi)", penalty_options)

                return params

            elif model_name == "Rasgele Orman":
                params["criterion"]= st.sidebar.selectbox("Criterion", ["gini", "entropy", "log_loss"])
                params["n_estimators"] = st.sidebar.slider("Ağaç Sayısı", 10, 300, 100)
                params["max_depth"] = st.sidebar.slider("Maksimum Derinlik", 1, 20, 10)


            elif model_name == "SVM":
                params["C"] = st.sidebar.slider("C (Regularization Strength)", 0.01, 10.0, 1.0)
                params["kernel"] = st.sidebar.selectbox("Kernel Tipi", ["linear", "poly", "rbf", "sigmoid"])

            elif model_name == "K-En Yakın Komşu":
                params["n_neighbors"] = st.sidebar.slider("Komşu Sayısı", 1, 15, 5)
                params["algorithm"] = st.sidebar.selectbox("Algoritma", ["auto", "ball_tree", "kd_tree", "brute"])
            return params

        # Model tanımlama
        def get_model(model_name, params):
            if model_name == "Lojistik Regresyon":
                model = LogisticRegression(solver= params['solver'],penalty=params["penalty"], max_iter=params["max_iter"], C=params["C"])
            elif model_name == "Rasgele Orman":
                model = RandomForestClassifier(n_estimators=params["n_estimators"], max_depth=params["max_depth"],criterion=params["criterion"])
            elif model_name == "SVM":
                model = SVC(C=params["C"], kernel=params["kernel"])
            elif model_name == "K-En Yakın Komşu":
                model = KNeighborsClassifier(n_neighbors=params["n_neighbors"], algorithm=params["algorithm"])
            return model

        # Modeli seçilen parametrelerle oluşturma
        params = model_hiperparametreleri(model_name)
        model = get_model(model_name, params)



        # Model eğitimi
        if st.button("Modeli Eğit"):
            model.fit(X_train, y_train)
            y_pred = model.predict(X_test)

            # Metrikleri hesapla
            accuracy = accuracy_score(y_test, y_pred)
            precision = precision_score(y_test, y_pred, average='weighted')
            recall = recall_score(y_test, y_pred, average='weighted')
            f1 = f1_score(y_test, y_pred, average='weighted')

            # Sonuçları göster
            st.write(f"Seçilen Model: {model_name}")
            st.write("Doğruluk (Accuracy):", accuracy)
            st.write("Kesinlik (Precision):", precision)
            st.write("Duyarlılık (Recall):", recall)
            st.write("F1 Skoru:", f1)

            # Feature importances ve feature_names'ı alıyoruz:
            if hasattr(model, 'feature_importances_'):
                feature_importances = model.feature_importances_
                feature_names = X_train.columns
            else:
                feature_importances = None
                feature_names = None


            # Raporu oluştur
            rapor_path = create_rapor(
                model_name=model_name,
                X_train=X_train,
                X_test=X_test,
                y_train=y_train,
                y_test=y_test,
                y_pred=y_pred,
                params=model.get_params() if hasattr(model, 'get_params') else None,
                feature_importances=feature_importances,
                feature_names=feature_names
            )

            # Kullanıcıya raporu indirme imkanı ver
            with open(rapor_path, "rb") as f:
                st.download_button("Raporu İndir", f, file_name=f'{model_name}_raporu_{datetime.now().strftime("%Y%m%d")}.docx')

            # Görselleştirme sekmeleri
            tab1, tab2, tab3, tab4, tab5 = st.tabs(["Confusion Matrix", "ROC-AUC Eğrisi", "Özellik Önem Düzeyleri", "Classification Report", "Hiper Parametreler"])

            with tab1:
                st.subheader("Confusion Matrix")
                cm = confusion_matrix(y_test, y_pred,)
                fig, ax = plt.subplots()
                sns.heatmap(cm, annot=True, fmt="d", cmap="Blues", cbar=False, ax=ax)
                st.pyplot(fig)
                

            with tab2:
                st.subheader("ROC-AUC Eğrisi")
                if hasattr(model, "predict_proba"):
                    y_prob = model.predict_proba(X_test)[:, 1]
                    fpr, tpr, _ = roc_curve(y_test, y_prob)
                    roc_auc = auc(fpr, tpr)
                    fig, ax = plt.subplots()
                    ax.plot(fpr, tpr, label=f'ROC Curve (AUC = {roc_auc:.2f})')
                    ax.plot([0, 1], [0, 1], 'k--')
                    ax.set_xlabel("False Positive Rate")
                    ax.set_ylabel("True Positive Rate")
                    ax.legend()
                    st.pyplot(fig)

            with tab3:
                st.subheader("Özellik Önem Düzeyleri")
                if model_name == "Rasgele Orman":
                    feature_importances = model.feature_importances_
                    importance_df = pd.DataFrame({"Özellik": X.columns, "Önem Düzeyi": feature_importances})
                    importance_df = importance_df.sort_values(by="Önem Düzeyi",ascending=False)
                    st.write(importance_df)
                    fig, ax = plt.subplots()
                    sns.barplot(data=importance_df, x="Önem Düzeyi", y="Özellik", ax=ax)
                    st.pyplot(fig)
                else:
                    st.write("Model Seçiminizde Bu Özellik Kullanılamıyor")
            
            with tab4:
                st.write("Sınıflandırma Raporu:")
                report = classification_report(y_test, y_pred,output_dict=True)
                report_df=pd.DataFrame(report).transpose()
                #fig, ax = plt.subplots()
                st.write(report_df)
                #st.text(report)
            
            with tab5:
                params = model.get_params()
                st.title(f"{model_name} Hiperparametreleri")
                param_df = pd.DataFrame(list(params.items()), columns=["Parametre", "Değer"])
                st.table(param_df)


        
            




        
           


           

        